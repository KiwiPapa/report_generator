# -*- coding: utf-8 -*-
import os, sys, time, xlrd, openpyxl, shutil
import easygui as gui
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from changeOffice import Change


# 转换文件，有问题，目前采用wps手动转换方法
# import win32com.client as wc
# word = wc.Dispatch("Word.Application")
# doc = word.Documents.Open(r"D:\#Programming Lab\Python_TestLab\\高石001-X43_20200217_原始资料收集登记表.doc")
# 上面的地方只能使用完整绝对地址，相对地址找不到文件，且，只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符。
# doc.SaveAs(r"D:\#Programming Lab\Python_TestLab\\test.docx", 16, False, "", True, "", False, \
# False, False, False)  # 转换后的文件
# doc.Close
# word.Quit

################################################################################
# 程序日志记录
class Logger(object):
    def __init__(self, filename='default.log', stream=sys.stdout):
        self.terminal = stream
        self.log = open(filename, 'w')

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)

    def flush(self):
        pass


# sys.stdout = Logger('.\\程序输出日志.log', sys.stdout)
# sys.stderr = Logger('.\\程序Bug日志.log_file', sys.stderr)
################################################################################
# 函数定义集结地

def mkdir(path):
    path = path.strip()  # 去除首位空格
    path = path.rstrip("\\")  # 去除尾部 \ 符号
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path)
        # print(path + ' 创建成功')
        return True
    else:
        # print(path + ' 目录已存在')
        return False


# 定义一个函数，增加重新计算后的厚度列
def get_thickness(x):
    thickness = x['井段End'] - x['井段Start']
    return thickness


# 定义进度条函数，用作进度展示
def view_bar(num, total):
    rate = float(num) / float(total)
    rate_num = int(rate * 100)
    r = '\r[%s%s]%d%%' % ("#" * rate_num, " " * (100 - rate_num), rate_num)
    sys.stdout.write(r)
    sys.stdout.flush()


# 函数，获取文件路径、文件名、后缀名
def get_filePath_fileName_fileExt(filename):
    (filepath, tempfilename) = os.path.split(filename)
    (shotname, extension) = os.path.splitext(tempfilename)
    return filepath, shotname, extension


# 文档替换主程序
def document_replace():
    newFile = PATH + "\\" + well_Name + '_' + year + month + \
              day + '_(' + casing_Goal + 'mm套,' + process_Section1 + 'm)固井报告' + '.docx'
    if formation_be_or_not == '有储层':
        document = Document(TEMPLATE_PATH + '\\template-with-formation.docx')
    else:
        document = Document(TEMPLATE_PATH + '\\template-without-formation.docx')
    document = check(document)
    # 全文档表格内容居中
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 整体设置，未起作用
        # table.style.font.color.rgb = RGBColor(255, 0, 0)
        # table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.save(newFile)


def check(document):
    # tables
    for table in document.tables:
        for row in range(len(table.rows)):
            view_bar(row, len(table.rows) - 1)  # 进度条
            for col in range(len(table.columns)):
                for key, value in DICT.items():
                    if key in table.cell(row, col).text:
                        # print(key + " = " + value)
                        table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

    # paragraphs
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            # view_bar(i, len(para.runs) - 1)
            for key, value in DICT.items():
                if key in para.runs[i].text:
                    # print(key + " = " + value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)

    # sections
    for sec in document.sections:
        for i in range(len(sec.header.paragraphs)):
            for key, value in DICT.items():
                if key in sec.header.paragraphs[i].text:
                    # print(key + " = " + value)
                    sec.header.paragraphs[i].text = sec.header.paragraphs[i].text.replace(key, value)
    return document

################################################################################
# 一界面某井段评价函数
def layer_evaluation1(df, start, end):
    df1 = df
    formation_Start = start
    formation_End = end
    df_temp = df1.loc[(df1['井段Start'] >= formation_Start) & (df1['井段Start'] <= formation_End), :]
    # 获取起始深度到第一层井段底界的结论
    df_temp_start_to_first_layer = df1.loc[(df1['井段Start'] <= formation_Start), :]
    start_to_upper_result = df_temp_start_to_first_layer.loc[len(df_temp_start_to_first_layer) - 1, '结论']
    # 获取calculation_Start所在段的声幅值
    df_temp_formation_Start = df1.loc[(df1['井段Start'] <= formation_Start) & (
            df1['井段End'] >= formation_Start), :]
    df_temp_formation_Start.reset_index(drop=True, inplace=True)  # 重新设置列索引#防止若截取中段，index不从0开始的bug
    # 补充储层界到井段的深度
    x, y = df_temp.shape
    df_temp = df_temp.reset_index()
    df_temp.drop(['index'], axis=1, inplace=True)
    if x != 0:  # 防止df_temp为空时，loc报错的bug
        first_layer_start = df_temp.loc[0, '井段Start']
    if x > 1 and first_layer_start != formation_Start:
        upper = pd.DataFrame({'井 段\n (m)': ''.join([str(formation_Start), '-', str(first_layer_start)]),
                              '厚 度\n (m)': first_layer_start - formation_Start,
                              '最大声幅\n （%）': df_temp_formation_Start.loc[0, '最大声幅\n （%）'],
                              '最小声幅\n  (%)': df_temp_formation_Start.loc[0, '最小声幅\n  (%)'],
                              '平均声幅\n  （%）': df_temp_formation_Start.loc[0, '平均声幅\n  （%）'],
                              '结论': start_to_upper_result,
                              '井段Start': formation_Start,
                              '井段End': first_layer_start},
                             index=[1])  # 自定义索引为：1 ，这里也可以不设置index
        df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
        df_temp = pd.concat([upper, df_temp], ignore_index=True)
        # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
        df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
        # 修改df_temp的最末一行
        df_temp.loc[len(df_temp) - 1, '井 段\n (m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                              '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
        df_temp.loc[len(df_temp) - 1, '厚 度\n (m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
    elif x > 1 and first_layer_start == formation_Start:
        df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
        df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
        # 修改df_temp的最末一行
        df_temp.loc[len(df_temp) - 1, '井 段\n (m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                              '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
        df_temp.loc[len(df_temp) - 1, '厚 度\n (m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
    else:  # 储层包含在一个井段内的情况
        df_temp = pd.DataFrame({'井 段\n (m)': ''.join([str(formation_Start), '-', str(formation_End)]),
                                '厚 度\n (m)': formation_End - formation_Start,
                                '最大声幅\n （%）': df_temp_formation_Start.loc[0, '最大声幅\n （%）'],
                                '最小声幅\n  (%)': df_temp_formation_Start.loc[0, '最小声幅\n  (%)'],
                                '平均声幅\n  （%）': df_temp_formation_Start.loc[0, '平均声幅\n  （%）'],
                                '结论': start_to_upper_result,
                                '井段Start': formation_Start,
                                '井段End': formation_End},
                               index=[1])  # 自定义索引为：1 ，这里也可以不设置index
        df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
        # 修改df_temp的最末一行
        df_temp.loc[len(df_temp), '井 段\n (m)'] = ''.join([str(df_temp.loc[len(df_temp), '井段Start']),
                                                          '-', str(df_temp.loc[len(df_temp), '井段End'])])
        df_temp.loc[len(df_temp), '厚 度\n (m)'] = df_temp.loc[len(df_temp), '重计算厚度']
    # print(df_temp)
    ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100
    if ratio_Series.__len__() == 2:
        if '好' not in ratio_Series:
            ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
        elif '中' not in ratio_Series:
            ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
        elif '差' not in ratio_Series:
            ratio_Series = ratio_Series.append(pd.Series({'差': 0}))
    elif ratio_Series.__len__() == 1:
        if ('好' not in ratio_Series) & ('中' not in ratio_Series):
            ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
            ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
        elif ('好' not in ratio_Series) & ('差' not in ratio_Series):
            ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
            ratio_Series = ratio_Series.append(pd.Series({'差': 0}))
        elif ('中' not in ratio_Series) & ('差' not in ratio_Series):
            ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
            ratio_Series = ratio_Series.append(pd.Series({'差': 0}))

    # 条件判断，参数需要研究
    if ratio_Series['好'] >= 95:
        evaluation_of_formation = '好'
    elif ratio_Series['中'] >= 95:
        evaluation_of_formation = '中'
    elif ratio_Series['差'] >= 95:
        evaluation_of_formation = '差'
    elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['中'] >= 5) & (5 >= ratio_Series['差']):
        if ratio_Series['好'] >= ratio_Series['中']:
            evaluation_of_formation = '中到好，以好为主'
        elif ratio_Series['好'] <= ratio_Series['中']:
            evaluation_of_formation = '中到好，以中等为主'
    elif (95 >= ratio_Series['差'] >= 5) & (95 >= ratio_Series['中'] >= 5) & (5 >= ratio_Series['好']):
        if ratio_Series['差'] >= ratio_Series['中']:
            evaluation_of_formation = '中到差，以差为主'
        elif ratio_Series['差'] <= ratio_Series['中']:
            evaluation_of_formation = '中到差，以中等为主'
    elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['差'] >= 5) & (5 >= ratio_Series['中']):
        if ratio_Series['好'] >= ratio_Series['差']:
            evaluation_of_formation = '好到差，以好为主'
        elif ratio_Series['好'] <= ratio_Series['差']:
            evaluation_of_formation = '好到差，以差为主'
    elif (80 >= ratio_Series['好'] >= 20) & (80 >= ratio_Series['差'] >= 20) & (80 >= ratio_Series['中'] >= 20):
        evaluation_of_formation = '好到中到差'
    print(ratio_Series)  # 某一储层的评价
    print(evaluation_of_formation)
    return ratio_Series, evaluation_of_formation

################################################################################
# 二界面某井段评价函数
def layer_evaluation2(df, start, end):
    df1 = df
    formation_Start = start
    formation_End = end
    df_temp = df1.loc[(df1['井段Start'] >= formation_Start) & (df1['井段Start'] <= formation_End), :]
    # 获取起始深度到第一层井段底界的结论
    df_temp_start_to_first_layer = df1.loc[(df1['井段Start'] <= formation_Start), :]
    start_to_upper_result = df_temp_start_to_first_layer.loc[len(df_temp_start_to_first_layer) - 1, '结论']
    # 获取calculation_Start所在段的声幅值
    df_temp_formation_Start = df1.loc[(df1['井段Start'] <= formation_Start) & (
            df1['井段End'] >= formation_Start), :]
    df_temp_formation_Start.reset_index(drop=True, inplace=True)  # 重新设置列索引#防止若截取中段，index不从0开始的bug
    # 补充储层界到井段的深度
    x, y = df_temp.shape
    df_temp = df_temp.reset_index()
    df_temp.drop(['index'], axis=1, inplace=True)
    if x != 0:  # 防止df_temp为空时，loc报错的bug
        first_layer_start = df_temp.loc[0, '井段Start']
    if x > 1 and first_layer_start != formation_Start:
        upper = pd.DataFrame({'井 段\n (m)': ''.join([str(formation_Start), '-', str(first_layer_start)]),
                              '厚 度\n (m)': first_layer_start - formation_Start,
                              '最大指数': df_temp_formation_Start.loc[0, '最大指数'],
                              '最小指数': df_temp_formation_Start.loc[0, '最小指数'],
                              '平均指数': df_temp_formation_Start.loc[0, '平均指数'],
                              '结论': start_to_upper_result,
                              '井段Start': formation_Start,
                              '井段End': first_layer_start},
                             index=[1])  # 自定义索引为：1 ，这里也可以不设置index
        df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
        df_temp = pd.concat([upper, df_temp], ignore_index=True)
        # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
        df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
        # 修改df_temp的最末一行
        df_temp.loc[len(df_temp) - 1, '井 段\n (m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                              '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
        df_temp.loc[len(df_temp) - 1, '厚 度\n (m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
    elif x > 1 and first_layer_start == formation_Start:
        df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
        df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
        # 修改df_temp的最末一行
        df_temp.loc[len(df_temp) - 1, '井 段\n (m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                              '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
        df_temp.loc[len(df_temp) - 1, '厚 度\n (m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
    else:  # 储层包含在一个井段内的情况
        df_temp = pd.DataFrame({'井 段\n (m)': ''.join([str(formation_Start), '-', str(formation_End)]),
                                '厚 度\n (m)': formation_End - formation_Start,
                                '最大指数': df_temp_formation_Start.loc[0, '最大指数'],
                                '最小指数': df_temp_formation_Start.loc[0, '最小指数'],
                                '平均指数': df_temp_formation_Start.loc[0, '平均指数'],
                                '结论': start_to_upper_result,
                                '井段Start': formation_Start,
                                '井段End': formation_End},
                               index=[1])  # 自定义索引为：1 ，这里也可以不设置index
        df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
        # 修改df_temp的最末一行
        df_temp.loc[len(df_temp), '井 段\n (m)'] = ''.join([str(df_temp.loc[len(df_temp), '井段Start']),
                                                          '-', str(df_temp.loc[len(df_temp), '井段End'])])
        df_temp.loc[len(df_temp), '厚 度\n (m)'] = df_temp.loc[len(df_temp), '重计算厚度']
    # print(df_temp)
    ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100
    if ratio_Series.__len__() == 2:
        if '好' not in ratio_Series:
            ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
        elif '中' not in ratio_Series:
            ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
        elif '差' not in ratio_Series:
            ratio_Series = ratio_Series.append(pd.Series({'差': 0}))
    elif ratio_Series.__len__() == 1:
        if ('好' not in ratio_Series) & ('中' not in ratio_Series):
            ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
            ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
        elif ('好' not in ratio_Series) & ('差' not in ratio_Series):
            ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
            ratio_Series = ratio_Series.append(pd.Series({'差': 0}))
        elif ('中' not in ratio_Series) & ('差' not in ratio_Series):
            ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
            ratio_Series = ratio_Series.append(pd.Series({'差': 0}))

    # 条件判断，参数需要研究
    if ratio_Series['好'] >= 95:
        evaluation_of_formation = '好'
    elif ratio_Series['中'] >= 95:
        evaluation_of_formation = '中'
    elif ratio_Series['差'] >= 95:
        evaluation_of_formation = '差'
    elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['中'] >= 5) & (5 >= ratio_Series['差']):
        if ratio_Series['好'] >= ratio_Series['中']:
            evaluation_of_formation = '中到好，以好为主'
        elif ratio_Series['好'] <= ratio_Series['中']:
            evaluation_of_formation = '中到好，以中等为主'
    elif (95 >= ratio_Series['差'] >= 5) & (95 >= ratio_Series['中'] >= 5) & (5 >= ratio_Series['好']):
        if ratio_Series['差'] >= ratio_Series['中']:
            evaluation_of_formation = '中到差，以差为主'
        elif ratio_Series['差'] <= ratio_Series['中']:
            evaluation_of_formation = '中到差，以中等为主'
    elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['差'] >= 5) & (5 >= ratio_Series['中']):
        if ratio_Series['好'] >= ratio_Series['差']:
            evaluation_of_formation = '好到差，以好为主'
        elif ratio_Series['好'] <= ratio_Series['差']:
            evaluation_of_formation = '好到差，以差为主'
    elif (80 >= ratio_Series['好'] >= 20) & (80 >= ratio_Series['差'] >= 20) & (80 >= ratio_Series['中'] >= 20):
        evaluation_of_formation = '好到中到差'
    print(ratio_Series)  # 某一储层的评价
    print(evaluation_of_formation)  # 全部储层的描述list
    return ratio_Series, evaluation_of_formation

################################################################################
# 转换文件，可能转出的文件读写空值，那么还得利用WPS或者LIBRE OFFICE


mkdir('.\\1原始资料-临时')
mkdir('.\\2统计表-临时')
mkdir('.\\3储层表-临时')

for fileName in os.listdir('.\\1原始资料'):
    if 'doc' in fileName and '~' not in fileName:
        shutil.copy('.\\1原始资料\\' + fileName, '.\\1原始资料-临时')
        pass
for fileName in os.listdir('.\\2统计表'):
    if 'xls' in fileName and '~' not in fileName:
        shutil.copy('.\\2统计表\\' + fileName, '.\\2统计表-临时')
        pass
for fileName in os.listdir('.\\3储层表'):
    if 'xls' in fileName and '~' not in fileName:
        shutil.copy('.\\3储层表\\' + fileName, '.\\3储层表-临时')
        pass
Goal_dir = Change(".\\1原始资料")
Goal_dir.doc2docx()
Goal_dir.xls2xlsx()
Goal_dir.ppt2pptx()
Goal_dir = Change(".\\2统计表")
# Goal_dir.doc2docx()
# Goal_dir.xls2xlsx()
# Goal_dir.ppt2pptx()
Goal_dir = Change(".\\3储层表")
Goal_dir.doc2docx()
Goal_dir.xls2xlsx()
Goal_dir.ppt2pptx()
for fileName in os.listdir('.\\1原始资料-临时'):
    if 'doc' in fileName:
        shutil.copy('.\\1原始资料-临时\\' + fileName, '.\\1原始资料')
        pass
for fileName in os.listdir('.\\2统计表-临时'):
    if 'doc' in fileName:
        shutil.copy('.\\2统计表-临时\\' + fileName, '.\\2统计表')
        pass
for fileName in os.listdir('.\\3储层表-临时'):
    if 'doc' in fileName:
        shutil.copy('.\\3储层表-临时\\' + fileName, '.\\3储层表')
        pass

isExists = os.path.exists('.\\1原始资料-临时')
if isExists:
    shutil.rmtree('.\\1原始资料-临时')
isExists = os.path.exists('.\\2统计表-临时')
if isExists:
    shutil.rmtree('.\\2统计表-临时')
isExists = os.path.exists('.\\3储层表-临时')
if isExists:
    shutil.rmtree('.\\3储层表-临时')
################################################################################
# 解析docx
PATH = ".\\1原始资料"
for fileName in os.listdir(PATH):
    if '.docx' in fileName and '~' not in fileName:
        fileDir = PATH + "\\" + fileName
    document = Document(fileDir)
'''
# 打印所有段落
print('一共有', str(len(document.paragraphs) - 1), '个文本段落。')
count = 0
for paragraph in document.paragraphs:
    print('第', str(count), '个段落')
    # print(paragraph.text)  # 打印各段落内容文本
    count += 1

# 打印所有表格
print('一共有', str(len(document.tables) - 1), '个表。')
count = 0
for table in document.tables:
    print('第', str(count), '个表')
    count += 1
    for row in range(len(table.rows)):
        for col in range(len(table.columns)):
            # table.cell(row, col).text += '({0},{1})'.format(row, col)#给文本中的单元格添加表格坐标
            print('(', str(row), ',', str(col), '):', table.cell(row, col).text)
'''
################################################################################
# 提取关键参数
well_Name_Raw = document.tables[0].cell(1, 2).text
well_Name = well_Name_Raw.split('井')
well_Name = well_Name[0]

well_Type = document.tables[0].cell(3, 2).text  # 井型
drilling_Unit = document.tables[0].cell(15, 2).text  # 钻井单位

bit1_Diameter = document.tables[0].cell(20, 2).text.strip()
bit1_Diameter = bit1_Diameter.replace(' ', '')
bit1_Diameter = bit1_Diameter.split('mm')
bit1_Diameter = bit1_Diameter[0]
bit1_Depth = document.tables[0].cell(20, 5).text.strip()
bit1_Depth = bit1_Depth.replace(' ', '')
bit1_Depth = bit1_Depth.split('m')
bit1_Depth = bit1_Depth[0]

bit2_Diameter = document.tables[0].cell(21, 2).text.strip()
bit2_Diameter = bit2_Diameter.replace(' ', '')
bit2_Diameter = bit2_Diameter.split('mm')
bit2_Diameter = bit2_Diameter[0]
bit2_Depth = document.tables[0].cell(21, 5).text.strip()
bit2_Depth = bit2_Depth.replace(' ', '')
bit2_Depth = bit2_Depth.split('m')
bit2_Depth = bit2_Depth[0]

bit3_Diameter = document.tables[0].cell(22, 2).text.strip()
bit3_Diameter = bit3_Diameter.replace(' ', '')
bit3_Diameter = bit3_Diameter.split('mm')
bit3_Diameter = bit3_Diameter[0]
bit3_Depth = document.tables[0].cell(22, 5).text.strip()
bit3_Depth = bit3_Depth.replace(' ', '')
bit3_Depth = bit3_Depth.split('m')
bit3_Depth = bit3_Depth[0]

bit4_Diameter = document.tables[0].cell(23, 2).text.strip()
bit4_Diameter = bit4_Diameter.replace(' ', '')
bit4_Diameter = bit4_Diameter.split('mm')
bit4_Diameter = bit4_Diameter[0]
bit4_Depth = document.tables[0].cell(23, 5).text.strip()
bit4_Depth = bit4_Depth.replace(' ', '')
bit4_Depth = bit4_Depth.split('m')
bit4_Depth = bit4_Depth[0]

bit5_Diameter = document.tables[0].cell(24, 2).text.strip()
bit5_Diameter = bit5_Diameter.replace(' ', '')
bit5_Diameter = bit5_Diameter.split('mm')
bit5_Diameter = bit5_Diameter[0]
bit5_Depth = document.tables[0].cell(24, 5).text.strip()
bit5_Depth = bit5_Depth.replace(' ', '')
bit5_Depth = bit5_Depth.split('m')
bit5_Depth = bit5_Depth[0]

bit6_Diameter = document.tables[0].cell(25, 2).text.strip()
bit6_Diameter = bit6_Diameter.replace(' ', '')
bit6_Diameter = bit6_Diameter.split('mm')
bit6_Diameter = bit6_Diameter[0]
bit6_Depth = document.tables[0].cell(25, 5).text.strip()
bit6_Depth = bit6_Depth.replace(' ', '')
bit6_Depth = bit6_Depth.split('m')
bit6_Depth = bit6_Depth[0]

bit7_Diameter = document.tables[0].cell(26, 2).text.strip()
bit7_Diameter = bit7_Diameter.replace(' ', '')
bit7_Diameter = bit7_Diameter.split('mm')
bit7_Diameter = bit7_Diameter[0]
bit7_Depth = document.tables[0].cell(26, 5).text.strip()
bit7_Depth = bit7_Depth.replace(' ', '')
bit7_Depth = bit7_Depth.split('m')
bit7_Depth = bit7_Depth[0]

bit8_Diameter = document.tables[0].cell(27, 2).text.strip()
bit8_Diameter = bit8_Diameter.replace(' ', '')
bit8_Diameter = bit8_Diameter.split('mm')
bit8_Diameter = bit8_Diameter[0]
bit8_Depth = document.tables[0].cell(27, 5).text.strip()
bit8_Depth = bit8_Depth.replace(' ', '')
bit8_Depth = bit8_Depth.split('m')
bit8_Depth = bit8_Depth[0]

bit9_Diameter = document.tables[0].cell(28, 2).text.strip()
bit9_Diameter = bit9_Diameter.replace(' ', '')
bit9_Diameter = bit9_Diameter.split('mm')
bit9_Diameter = bit9_Diameter[0]
bit9_Depth = document.tables[0].cell(28, 5).text.strip()
bit9_Depth = bit9_Depth.replace(' ', '')
bit9_Depth = bit9_Depth.split('m')
bit9_Depth = bit9_Depth[0]

bit10_Diameter = document.tables[0].cell(29, 2).text.strip()
bit10_Diameter = bit10_Diameter.replace(' ', '')
bit10_Diameter = bit10_Diameter.split('mm')
bit10_Diameter = bit10_Diameter[0]
bit10_Depth = document.tables[0].cell(29, 5).text.strip()
bit10_Depth = bit10_Depth.replace(' ', '')
bit10_Depth = bit10_Depth.split('m')
bit10_Depth = bit10_Depth[0]

# 找出最深的钻头深度deepest_bit
if bit10_Depth != '':
    deepest_bit = bit10_Depth
elif bit9_Depth != '':
    deepest_bit = bit9_Depth
elif bit8_Depth != '':
    deepest_bit = bit8_Depth
elif bit7_Depth != '':
    deepest_bit = bit7_Depth
elif bit6_Depth != '':
    deepest_bit = bit6_Depth
elif bit5_Depth != '':
    deepest_bit = bit5_Depth
elif bit4_Depth != '':
    deepest_bit = bit4_Depth
elif bit3_Depth != '':
    deepest_bit = bit3_Depth
elif bit2_Depth != '':
    deepest_bit = bit2_Depth
elif bit1_Depth != '':
    deepest_bit = bit1_Depth

# 地理位置geo_Position
geographic_Position = document.tables[0].cell(30, 2).text.strip()
if geographic_Position != '':
    if '省' in geographic_Position:
        geographic_Position = geographic_Position.split('省')
        geographic_Position1 = ''.join([geographic_Position[0], '省'])
        if '县' not in geographic_Position[1]:
            geographic_Position2 = geographic_Position[1].split('市')
            geographic_Position2 = ''.join([geographic_Position2[0], '市'])
        else:
            geographic_Position2 = geographic_Position[1].split('县')
            geographic_Position2 = geographic_Position2[0]
            if '市' in geographic_Position2:
                geographic_Position2 = geographic_Position2.split('市')[1]
            geographic_Position2 = ''.join([geographic_Position2, '县'])
    elif '省' not in geographic_Position:
        geographic_Position = geographic_Position.split('市')
        geographic_Position1 = ''.join([geographic_Position[0], '市'])
        geographic_Position2 = geographic_Position[1].split('区')
        geographic_Position2 = ''.join([geographic_Position2[0], '区'])
    geo_Position = ''.join([geographic_Position1, geographic_Position2])

# 构造位置stru_Position
structure_Position = document.tables[0].cell(31, 2).text
if structure_Position != '':
    structure_Position = structure_Position.replace(' ', '')
    structure_Position = structure_Position.replace('四川盆地', '')
    structure_Position = structure_Position.split('构造')
    stru_Position = structure_Position[0]
    if '高石' in stru_Position:
        stru_Position = '高石梯'
    elif '磨溪' in stru_Position:
        stru_Position = '磨溪'
    elif '威远' in stru_Position:
        stru_Position = '威远'

# 钻井液flu_Property, flu_Density, flu_Viscosity
flu_Property = document.tables[1].cell(9, 2).text.strip()
flu_Property = flu_Property.replace(' ', '')

drilling_Fluid_Density = document.tables[1].cell(10, 2).text.strip()
drilling_Fluid_Density = drilling_Fluid_Density.replace(' ', '')
drilling_Fluid_Density = drilling_Fluid_Density.split('g')
flu_Density = drilling_Fluid_Density[0]

drilling_Fluid_Viscosity = document.tables[1].cell(11, 2).text.strip()
drilling_Fluid_Viscosity = drilling_Fluid_Viscosity.replace(' ', '')
if 's' in drilling_Fluid_Viscosity:
    drilling_Fluid_Viscosity = drilling_Fluid_Viscosity.split('s')
elif 'S' in drilling_Fluid_Viscosity:
    drilling_Fluid_Viscosity = drilling_Fluid_Viscosity.split('S')
elif '秒' in drilling_Fluid_Viscosity:
    drilling_Fluid_Viscosity = drilling_Fluid_Viscosity.split('秒')
flu_Viscosity = drilling_Fluid_Viscosity[0]

# 测井装备
logging_Equipment = document.tables[1].cell(13, 1).text.strip()
# 测井小队
logging_Group = document.tables[1].cell(14, 1).text.strip()
# 小队长
logging_Leader = document.tables[1].cell(15, 1).text.strip()

# 时间cement_End_Time, logging_Start_Time, logging_End_Time
cement_End_Time = document.tables[1].cell(14, 4).text.strip()
cement_End_Time = cement_End_Time[0:10]

logging_Start_Time = document.tables[1].cell(15, 4).text.strip()
logging_Start_Time = logging_Start_Time[0:10]

logging_End_Time = document.tables[1].cell(16, 4).text.strip()
log_End_Time = logging_End_Time[0:10]

year = logging_End_Time[0:4]
month = logging_End_Time[5:7]
day = logging_End_Time[8:10]

# 最大井斜
max_Well_Deviation = document.tables[2].cell(1, 2).text
max_Well_Deviation = max_Well_Deviation.replace(' ', '')
max_Well_Deviation = max_Well_Deviation.replace('°', '')

# 最大井斜深度
max_Well_Deviation_Depth = document.tables[2].cell(1, 7).text
max_Well_Deviation_Depth = max_Well_Deviation_Depth.replace(' ', '')
max_Well_Deviation_Depth = max_Well_Deviation_Depth.replace('m', '')

dev_Depth_Ratio = ''.join([max_Well_Deviation, '/', max_Well_Deviation_Depth])

# 人工井底arti_Bottom
artificial_Bottom_of_Well = document.tables[2].cell(2, 2).text.strip()
artificial_Bottom_of_Well = artificial_Bottom_of_Well.replace(' ', '')
artificial_Bottom_of_Well = artificial_Bottom_of_Well.replace('m', '')
if '.' in artificial_Bottom_of_Well:
    arti_Bottom = artificial_Bottom_of_Well
else:
    if artificial_Bottom_of_Well != '':
        arti_Bottom = ''.join([artificial_Bottom_of_Well, '.00'])
    else:
        arti_Bottom = ''

# 已注入水泥量cement_Quantity
cement_Quantity = document.tables[2].cell(3, 7).text
cement_Quantity = cement_Quantity.replace(' ', '')
cement_Quantity = cement_Quantity.replace('T', '')
cement_Quantity = cement_Quantity.replace('t', '')

# 水泥密度cement_Density
cement_Density = ''
slow_Cement_Density = document.tables[2].cell(7, 7).text.strip()
fast_Cement_Density = document.tables[2].cell(8, 7).text.strip()
if slow_Cement_Density == '':
    cement_Density = fast_Cement_Density
elif fast_Cement_Density == '':
    cement_Density = slow_Cement_Density
elif eval(str(slow_Cement_Density)) == eval(str(fast_Cement_Density)):
    cement_Density = fast_Cement_Density
elif eval(str(slow_Cement_Density)) > eval(str(fast_Cement_Density)):
    cement_Density = ''.join([fast_Cement_Density, '~', slow_Cement_Density])
elif eval(str(slow_Cement_Density)) < eval(str(fast_Cement_Density)):
    cement_Density = ''.join([slow_Cement_Density, '~', fast_Cement_Density])
other_Cement_Density = document.tables[2].cell(9, 7).text.strip()
if other_Cement_Density != '' and cement_Density == '':
    cement_Density = other_Cement_Density

# 水泥设计返高design_Depth
design_Depth = document.tables[2].cell(5, 2).text.strip()
design_Depth = design_Depth.replace(' ', '')
design_Depth = design_Depth.replace('m', '')
if design_Depth == ['井口', '地面']:
    design_Depth = '0'
else:
    design_Depth = design_Depth

# 水泥实际返高actual_Depth
actual_Depth = document.tables[2].cell(5, 7).text.strip()
actual_Depth = actual_Depth.replace(' ', '')
actual_Depth = actual_Depth.replace('m', '')
actual_Depth = actual_Depth.replace('以上', '')
actual_Depth = actual_Depth.replace('（', '')
actual_Depth = actual_Depth.replace('）', '')
actual_Depth = actual_Depth.replace('(', '')
actual_Depth = actual_Depth.replace(')', '')
if '.' in actual_Depth:
    actual_Depth = actual_Depth
else:
    if actual_Depth != '':
        actual_Depth = ''.join([actual_Depth, '.00'])
    else:
        actual_Depth = ''

# 套管数据
casing1_Dia = document.tables[2].cell(15, 3).text.strip()
casing2_Dia = document.tables[2].cell(16, 3).text.strip()
casing3_Dia = document.tables[2].cell(17, 3).text.strip()
casing4_Dia = document.tables[2].cell(18, 3).text.strip()
casing5_Dia = document.tables[2].cell(19, 3).text.strip()
casing6_Dia = document.tables[2].cell(20, 3).text.strip()
casing7_Dia = document.tables[2].cell(21, 3).text.strip()
casing8_Dia = document.tables[2].cell(22, 3).text.strip()
casing9_Dia = document.tables[2].cell(23, 3).text.strip()
casing10_Dia = document.tables[2].cell(24, 3).text.strip()
casing11_Dia = document.tables[2].cell(25, 3).text.strip()
casing12_Dia = document.tables[2].cell(26, 3).text.strip()

# 避免套管下深井段为单数字而不为井段
casing1_interval = document.tables[2].cell(15, 6).text.strip()
casing1_interval = casing1_interval.replace(' ', '')
casing1_interval = casing1_interval.replace('～', '-')
casing1_interval = casing1_interval.replace('~', '-')
if '~' not in casing1_interval and '～' not in casing1_interval and \
        '-' not in casing1_interval and casing1_interval != '':
    casing1_interval = ''.join(['0', '~', casing1_interval])
casing2_interval = document.tables[2].cell(16, 6).text.strip()
casing2_interval = casing2_interval.replace(' ', '')
casing2_interval = casing2_interval.replace('～', '-')
casing2_interval = casing2_interval.replace('~', '-')
if '~' not in casing2_interval and '～' not in casing2_interval and \
        '-' not in casing2_interval and casing2_interval != '':
    casing2_interval = ''.join(['0', '~', casing2_interval])
casing3_interval = document.tables[2].cell(17, 6).text.strip()
casing3_interval = casing3_interval.replace(' ', '')
casing3_interval = casing3_interval.replace('～', '-')
casing3_interval = casing3_interval.replace('~', '-')
if '~' not in casing3_interval and '～' not in casing3_interval and \
        '-' not in casing3_interval and casing3_interval != '':
    casing3_interval = ''.join(['0', '~', casing3_interval])
casing4_interval = document.tables[2].cell(18, 6).text.strip()
casing4_interval = casing4_interval.replace(' ', '')
casing4_interval = casing4_interval.replace('～', '-')
casing4_interval = casing4_interval.replace('~', '-')
if '~' not in casing4_interval and '～' not in casing4_interval and \
        '-' not in casing4_interval and casing4_interval != '':
    casing4_interval = ''.join(['0', '~', casing4_interval])
casing5_interval = document.tables[2].cell(19, 6).text.strip()
casing5_interval = casing5_interval.replace(' ', '')
casing5_interval = casing5_interval.replace('～', '-')
casing5_interval = casing5_interval.replace('~', '-')
if '~' not in casing5_interval and '～' not in casing5_interval and \
        '-' not in casing5_interval and casing5_interval != '':
    casing5_interval = ''.join(['0', '~', casing5_interval])
casing6_interval = document.tables[2].cell(20, 6).text.strip()
casing6_interval = casing6_interval.replace(' ', '')
casing6_interval = casing6_interval.replace('～', '-')
casing6_interval = casing6_interval.replace('~', '-')
if '~' not in casing6_interval and '～' not in casing6_interval and \
        '-' not in casing6_interval and casing6_interval != '':
    casing6_interval = ''.join(['0', '~', casing6_interval])
casing7_interval = document.tables[2].cell(21, 6).text.strip()
casing7_interval = casing7_interval.replace(' ', '')
casing7_interval = casing7_interval.replace('～', '-')
casing7_interval = casing7_interval.replace('~', '-')
if '~' not in casing7_interval and '～' not in casing7_interval and \
        '-' not in casing7_interval and casing7_interval != '':
    casing7_interval = ''.join(['0', '~', casing7_interval])
casing8_interval = document.tables[2].cell(22, 6).text.strip()
casing8_interval = casing8_interval.replace(' ', '')
casing8_interval = casing8_interval.replace('～', '-')
casing8_interval = casing8_interval.replace('~', '-')
if '~' not in casing8_interval and '～' not in casing8_interval and \
        '-' not in casing8_interval and casing8_interval != '':
    casing8_interval = ''.join(['0', '~', casing8_interval])
casing9_interval = document.tables[2].cell(23, 6).text.strip()
casing9_interval = casing9_interval.replace(' ', '')
casing9_interval = casing9_interval.replace('～', '-')
casing9_interval = casing9_interval.replace('~', '-')
if '~' not in casing9_interval and '～' not in casing9_interval and \
        '-' not in casing9_interval and casing9_interval != '':
    casing9_interval = ''.join(['0', '~', casing9_interval])
casing10_interval = document.tables[2].cell(24, 6).text.strip()
casing10_interval = casing10_interval.replace(' ', '')
casing10_interval = casing10_interval.replace('～', '-')
casing10_interval = casing10_interval.replace('~', '-')
if '~' not in casing10_interval and '～' not in casing10_interval and \
        '-' not in casing10_interval and casing10_interval != '':
    casing10_interval = ''.join(['0', '~', casing10_interval])
casing11_interval = document.tables[2].cell(25, 6).text.strip()
casing11_interval = casing11_interval.replace(' ', '')
casing11_interval = casing11_interval.replace('～', '-')
casing11_interval = casing11_interval.replace('~', '-')
if '~' not in casing11_interval and '～' not in casing11_interval and \
        '-' not in casing11_interval and casing11_interval != '':
    casing11_interval = ''.join(['0', '~', casing11_interval])
casing12_interval = document.tables[2].cell(26, 6).text.strip()
casing12_interval = casing12_interval.replace(' ', '')
casing12_interval = casing12_interval.replace('～', '-')
casing12_interval = casing12_interval.replace('~', '-')
if '~' not in casing12_interval and '～' not in casing12_interval and \
        '-' not in casing12_interval and casing12_interval != '':
    casing12_interval = ''.join(['0', '~', casing12_interval])

# 目标套管尺寸casing_Goal
if casing12_Dia != '':
    casing_Goal = casing12_Dia
elif casing11_Dia != '':
    casing_Goal = casing11_Dia
elif casing10_Dia != '':
    casing_Goal = casing10_Dia
elif casing9_Dia != '':
    casing_Goal = casing9_Dia
elif casing8_Dia != '':
    casing_Goal = casing8_Dia
elif casing7_Dia != '':
    casing_Goal = casing7_Dia
elif casing6_Dia != '':
    casing_Goal = casing6_Dia
elif casing5_Dia != '':
    casing_Goal = casing5_Dia
elif casing4_Dia != '':
    casing_Goal = casing4_Dia
elif casing3_Dia != '':
    casing_Goal = casing3_Dia
elif casing2_Dia != '':
    casing_Goal = casing2_Dia
elif casing1_Dia != '':
    casing_Goal = casing1_Dia

# 目标套管下深casing_Goal_Depth
if casing12_Dia == casing_Goal:
    if '～' in document.tables[2].cell(26, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(26, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(26, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(26, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(26, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(26, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(26, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(26, 6).text.strip()
elif casing11_Dia == casing_Goal:
    if '～' in document.tables[2].cell(25, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(25, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(25, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(25, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(25, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(25, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(25, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(25, 6).text.strip()
elif casing10_Dia == casing_Goal:
    if '～' in document.tables[2].cell(24, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(24, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(24, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(24, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(24, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(24, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(24, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(24, 6).text.strip()
elif casing9_Dia == casing_Goal:
    if '～' in document.tables[2].cell(23, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(23, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(23, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(23, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(23, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(23, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(23, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(23, 6).text.strip()
elif casing8_Dia == casing_Goal:
    if '～' in document.tables[2].cell(22, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(22, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(22, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(22, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(22, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(22, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(22, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(22, 6).text.strip()
elif casing7_Dia == casing_Goal:
    if '～' in document.tables[2].cell(21, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(21, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(21, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(21, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(21, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(21, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(21, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(21, 6).text.strip()
elif casing6_Dia == casing_Goal:
    if '～' in document.tables[2].cell(20, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(20, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(20, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(20, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(20, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(20, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(20, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(20, 6).text.strip()
elif casing5_Dia == casing_Goal:
    if '～' in document.tables[2].cell(19, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(19, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(19, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(19, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(19, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(19, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(19, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(19, 6).text.strip()
elif casing4_Dia == casing_Goal:
    if '～' in document.tables[2].cell(18, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(18, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(18, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(18, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(18, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(18, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(18, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(18, 6).text.strip()
elif casing3_Dia == casing_Goal:
    if '～' in document.tables[2].cell(17, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(17, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(17, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(17, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(17, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(17, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(17, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(17, 6).text.strip()
elif casing2_Dia == casing_Goal:
    if '～' in document.tables[2].cell(16, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(16, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(16, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(16, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(16, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(16, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(16, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(16, 6).text.strip()
elif casing1_Dia == casing_Goal:
    if '～' in document.tables[2].cell(15, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(15, 6).text.strip().split('～')[1]
    elif '~' in document.tables[2].cell(15, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(15, 6).text.strip().split('~')[1]
    elif '-' in document.tables[2].cell(15, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(15, 6).text.strip().split('-')[1]
    elif '-' not in document.tables[2].cell(15, 6).text.strip():
        casing_Goal_Depth = document.tables[2].cell(15, 6).text.strip()

# 获取测量井段
for row in range(3, 26):
    if document.tables[3].cell(row, 5).text.strip() != '':
        measure_Interval = document.tables[3].cell(row, 5).text.strip()
        measure_Interval = measure_Interval.replace('~', '-')
        measure_Interval = measure_Interval.replace('～', '-')
        pass

# 判断甲方是谁
if len(document.tables) == 9:
    if document.tables[8].cell(2, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(2, 2).text.strip()
    elif document.tables[8].cell(3, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(3, 2).text.strip()
    elif document.tables[8].cell(4, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(4, 2).text.strip()
    elif document.tables[8].cell(5, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(5, 2).text.strip()
    elif document.tables[8].cell(6, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(6, 2).text.strip()
    elif document.tables[8].cell(7, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(7, 2).text.strip()
    elif document.tables[8].cell(8, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(8, 2).text.strip()
    elif document.tables[8].cell(9, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(9, 2).text.strip()
    elif document.tables[8].cell(10, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(10, 2).text.strip()
    elif document.tables[8].cell(11, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(11, 2).text.strip()
    elif document.tables[8].cell(12, 6).text.strip() == '√':
        client_Name = document.tables[8].cell(12, 2).text.strip()
elif len(document.tables) == 8:
    if document.tables[7].cell(2, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(2, 2).text.strip()
    elif document.tables[7].cell(3, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(3, 2).text.strip()
    elif document.tables[7].cell(4, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(4, 2).text.strip()
    elif document.tables[7].cell(5, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(5, 2).text.strip()
    elif document.tables[7].cell(6, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(6, 2).text.strip()
    elif document.tables[7].cell(7, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(7, 2).text.strip()
    elif document.tables[7].cell(8, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(8, 2).text.strip()
    elif document.tables[7].cell(9, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(9, 2).text.strip()
    elif document.tables[7].cell(10, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(10, 2).text.strip()
    elif document.tables[7].cell(11, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(11, 2).text.strip()
    elif document.tables[7].cell(12, 6).text.strip() == '√':
        client_Name = document.tables[7].cell(12, 2).text.strip()

check_or_not = gui.indexbox(msg="是否进行【原始资料收集登记表】完整性检查？", title="提示", choices=("检查", "不检查"))

DICT_CHECK = {
    "井名": well_Name,
    "构造位置": stru_Position,
    "目标套管": casing_Goal,
    "测量井段": measure_Interval,
    "地理位置": geo_Position,
    "人工井底": arti_Bottom,
    "最大井斜深度": dev_Depth_Ratio,
    "套管1内径": casing1_Dia,
    "钻头1直径": bit1_Diameter,
    "流体性质": flu_Property,
    "流体密度": flu_Density,
    "流体粘度": flu_Viscosity,
    "水泥密度": cement_Density,
    "水泥量": cement_Quantity,
    "水泥设计返高": design_Depth,
    "固井结束时间": cement_End_Time,
    "测井结束时间": log_End_Time,
    "测井小队": logging_Group,
    "小队长": logging_Leader,
    "测井系列": logging_Equipment,
    "水泥实际返高": actual_Depth,
    "年": year,
    "月": month,
    "套管1井段": casing1_interval,
    "钻头1深度": bit1_Depth
}
remind1 = []
remind2 = []
remind = []
if check_or_not == 0:
    for k, v in DICT_CHECK.items():
        if v != '':
            temp_str1 = ''.join([k, '：', v, '\n'])
            remind1.append(temp_str1)
        elif v == '':
            temp_str2 = ''.join(['请注意:', k, '没有填写\n'])
            remind2.append(temp_str2)
    remind = ''.join(remind1 + remind2)
    if remind != '':
        if actual_Depth == '':
            remind.append('检查完毕，若没问题，可继续。')
            #remind.append('再次提醒：水泥返高一定要填写，若不填写可能会报错。')
        remind = ''.join(remind)
        gui.msgbox(msg=remind, title="提示")
    else:
        remind = '原始资料登记卡填写完整:)'
        gui.msgbox(msg=remind, title="提示")
    print('完整性检查完毕')

print('【原始资料收集登记表】解析完成')

# 现在开始提取成果表中的内容
PATH = ".\\2统计表"
for fileName in os.listdir(PATH):
    if '1统' in fileName:
        fileDir = PATH + "\\" + fileName
        workbook1 = xlrd.open_workbook(fileDir)
    elif '2统' in fileName:
        fileDir = PATH + "\\" + fileName
        workbook2 = xlrd.open_workbook(fileDir)

##########################
# 解析解释成果表-1统
sheet1 = workbook1.sheets()[0]

nrow1 = sheet1.nrows
ncol1 = sheet1.ncols

# 统计结论
good_Length1 = str(sheet1.cell_value(3, 2))
good_Ratio1 = str(sheet1.cell_value(3, 3))

median_Length1 = str(sheet1.cell_value(4, 2))
median_Ratio1 = str(sheet1.cell_value(4, 3))

bad_Length1 = str(sheet1.cell_value(5, 2))
bad_Ratio1 = str(sheet1.cell_value(5, 3))

# 合格率
pass_Percent1 = str(round((sheet1.cell_value(3, 3) + sheet1.cell_value(4, 3)), 2))

##########################
# 解析解释成果表-2统
sheet2 = workbook2.sheets()[0]

nrow2 = sheet2.nrows
ncol2 = sheet2.ncols

# 统计结论
good_Length2 = str(sheet2.cell_value(3, 2))
good_Ratio2 = str(sheet2.cell_value(3, 3))

median_Length2 = str(sheet2.cell_value(4, 2))
median_Ratio2 = str(sheet2.cell_value(4, 3))

bad_Length2 = str(sheet2.cell_value(5, 2))
bad_Ratio2 = str(sheet2.cell_value(5, 3))

# 合格率
pass_Percent2 = str(round((sheet2.cell_value(3, 3) + sheet2.cell_value(4, 3)), 2))
print('【统计表】解析完成')

# 整体评价
if eval(pass_Percent1) >= 60:
    eval_Result1 = '合格'
else:
    eval_Result1 = '不合格'

if eval(pass_Percent2) >= 60:
    eval_Result2 = '合格'
else:
    eval_Result2 = '不合格'
################################################################################
# 读取单层统计表
PATH = ".\\2统计表"
for fileName in os.listdir(PATH):
    if '1单' in fileName and '.xlsx' in fileName and '$' not in fileName:
        fileDir = PATH + "\\" + fileName
        workbook = xlrd.open_workbook(fileDir)

sheet = workbook.sheets()[0]

# 获得表单的行数及列数
nrow = sheet.nrows
ncol = sheet.ncols
# 处理评价井段
start_Evaluation = str(sheet.cell_value(4, 1)).strip()
start_Evaluation = start_Evaluation.split('-')[0]
end_Evaluation = str(sheet.cell_value(nrow - 1, 1)).strip('')
end_Evaluation = ''.join(end_Evaluation.split())  # 去除所有空格
end_Evaluation = end_Evaluation.split('-')[1]
process_Section1 = ''.join([start_Evaluation, '-', end_Evaluation])

PATH = ".\\2统计表"
for fileName in os.listdir(PATH):
    if '2单' in fileName and '.xlsx' in fileName and '$' not in fileName:
        fileDir = PATH + "\\" + fileName
        workbook = xlrd.open_workbook(fileDir)

sheet = workbook.sheets()[0]

# 获得表单的行数及列数
nrow = sheet.nrows
ncol = sheet.ncols
# 处理评价井段
start_Evaluation = str(sheet.cell_value(4, 1)).strip()
start_Evaluation = start_Evaluation.split('-')[0]
end_Evaluation = str(sheet.cell_value(nrow - 1, 1)).strip('')
end_Evaluation = ''.join(end_Evaluation.split())  # 去除所有空格
end_Evaluation = end_Evaluation.split('-')[1]
process_Section2 = ''.join([start_Evaluation, '-', end_Evaluation])

##########################
# 液面高度的获取
if gui.ccbox("请问液面高度是否和开始评价深度一致？", choices=('不一致', '一致')):
    fluid_Height = gui.enterbox(msg='请输入液面高度', title='提示')
else:
    fluid_Height = start_Evaluation
################################################################################
# 判断是否有储层
PATH = ".\\3储层表\\"
formation_be_or_not = ''
if os.listdir(PATH) != []:
    for fileName in os.listdir(PATH):
        fileName = fileName
else:
    fileName = ''

f_path = PATH + fileName

if os.path.isdir(f_path):
    formation_be_or_not = '无储层'
else:
    formation_be_or_not = '有储层'
################################################################################
# 储层表解析
if formation_be_or_not == '有储层':
    PATH = ".\\3储层表"
    for fileName in os.listdir(PATH):
        fileDir = PATH + "\\" + fileName
        workbook = xlrd.open_workbook(fileDir)

    sheet = workbook.sheets()[0]

    # 通过xlrd的接口获得表单的行数及列数
    nrow = sheet.nrows
    ncol = sheet.ncols

    if nrow >= 3:
        formation_Number = str(nrow - 2)
    else:
        formation_Number = '[待确定]'
    print('【储层表】解析完成')
else:
    print('未发现可供解析的储层表')
################################################################################
# 储层表和单层统计表的联动数据分析
all_evaluation_of_formation = []
if formation_be_or_not == '有储层':
    # 单层统计表
    PATH = ".\\2统计表"
    for fileName in os.listdir(PATH):
        if '1单' in fileName and '.xlsx' in fileName and '$' not in fileName:
            fileDir = PATH + "\\" + fileName
    df1 = pd.read_excel(fileDir, header=2, index='序号')
    df1.drop([0], inplace=True)
    df1.loc[:, '井 段\n (m)'] = df1['井 段\n (m)'].str.replace(' ', '') # 消除数据中空格
    df1['井段Start'] = df1['井 段\n (m)'].map(lambda x: x.split("-")[0])
    df1['井段End'] = df1['井 段\n (m)'].map(lambda x: x.split("-")[1])

    # 储层表
    PATH = ".\\3储层表"
    for fileName in os.listdir(PATH):
        fileDir = PATH + "\\" + fileName
    df2 = pd.read_excel(fileDir, header=0, index='序号')
    df2.drop([0], inplace=True)
    df2.drop(['层位', '解释结论'], axis=1, inplace=True)
    df2.loc[:, '井        段'] = df2['井        段'].str.replace(' ', '')  # 消除数据中空格
    df2['储层Start'] = df2['井        段'].map(lambda x: x.split("--")[0])
    df2['储层End'] = df2['井        段'].map(lambda x: x.split("--")[1])
    # print(df2)

    # 表格数据清洗
    df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
    df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
    df2.loc[:, "储层Start"] = df2["储层Start"].str.replace(" ", "").astype('float')
    df2.loc[:, "储层End"] = df2["储层End"].str.replace(" ", "").astype('float')
    rows1, cols1 = df1.shape
    rows2, cols2 = df2.shape

    # 针对每个储层在单层评价表中得出好中差比例
    for row in range(1, rows2 + 1):
        formation_Start = df2.loc[row, '储层Start']
        formation_End = df2.loc[row, '储层End']
        # print('----------------第', row, '个储层内的井段----------------')
        if (formation_End <= float(end_Evaluation)) & (formation_Start >= float(start_Evaluation)):
            ratio_Series = layer_evaluation1(df1, formation_Start, formation_End)[0] #调取一界面评价函数
            evaluation_of_formation = layer_evaluation1(df1, formation_Start, formation_End)[1] #调取一界面评价函数
            all_evaluation_of_formation.append(evaluation_of_formation)
        elif (formation_End > float(end_Evaluation)) & (formation_Start < float(end_Evaluation)) & (formation_Start >= float(start_Evaluation)):
            ratio_Series = layer_evaluation1(df1, formation_Start, end_Evaluation)[0]  # 调取一界面评价函数
            evaluation_of_formation = layer_evaluation1(df1, formation_Start, end_Evaluation)[1]  # 调取一界面评价函数
        else:
            print('储层界超出了测量范围，请检查')
            pass
# print(all_evaluation_of_formation)
################################################################################
# 基于文本替换方案的文档生成
TEMPLATE_PATH = ".\\resources\\模板报告"
PATH = "."

DICT = {
    "well_Name": well_Name,
    "stru_Position": stru_Position,
    "casing_Goal": casing_Goal,
    "start_Evaluation": start_Evaluation,
    "end_Evaluation": end_Evaluation,
    "measure_Interval": measure_Interval,
    "process_Section1": process_Section1,
    "process_Section2": process_Section2,
    "geo_Position": geo_Position,
    "deepest_bit": deepest_bit,
    "arti_Bottom": arti_Bottom,
    "dev_Depth_Ratio": dev_Depth_Ratio,
    "casing1_Dia": casing1_Dia,
    "casing2_Dia": casing2_Dia,
    "casing3_Dia": casing3_Dia,
    "casing4_Dia": casing4_Dia,
    "casing5_Dia": casing5_Dia,
    "bit1_Diameter": bit1_Diameter,
    "bit2_Diameter": bit2_Diameter,
    "bit3_Diameter": bit3_Diameter,
    "bit4_Diameter": bit4_Diameter,
    "bit5_Diameter": bit5_Diameter,
    "flu_Property": flu_Property,
    "flu_Density": flu_Density,
    "flu_Viscosity": flu_Viscosity,
    "cement_Density": cement_Density,
    "cement_Quantity": cement_Quantity,
    "design_Depth": design_Depth,
    "cement_End_Time": cement_End_Time,
    "log_End_Time": log_End_Time,
    "logging_Group": logging_Group,
    "logging_Leader": logging_Leader,
    "logging_Equipment": logging_Equipment,
    "actual_Depth": actual_Depth,
    "good_Ratio1": good_Ratio1,
    "median_Ratio1": median_Ratio1,
    "bad_Ratio1": bad_Ratio1,
    "pass_Percent1": pass_Percent1,
    "eval_Result1": eval_Result1,
    "good_Length1": good_Length1,
    "median_Length1": median_Length1,
    "bad_Length1": bad_Length1,
    "good_Ratio2": good_Ratio2,
    "median_Ratio2": median_Ratio2,
    "bad_Ratio2": bad_Ratio2,
    "pass_Percent2": pass_Percent2,
    "eval_Result2": eval_Result2,
    "good_Length2": good_Length2,
    "median_Length2": median_Length2,
    "bad_Length2": bad_Length2,
    "year": year,
    "month": month,
    "casing1_interval": casing1_interval,
    "casing2_interval": casing2_interval,
    "casing3_interval": casing3_interval,
    "casing4_interval": casing4_interval,
    "casing5_interval": casing5_interval,
    "bit1_Depth": bit1_Depth,
    "bit2_Depth": bit2_Depth,
    "bit3_Depth": bit3_Depth,
    "bit4_Depth": bit4_Depth,
    "bit5_Depth": bit5_Depth,
    "fluid_Height": fluid_Height
}

print('模板替换开始，请等待...')
document_replace()  # 模板替换主程序
print('\n【模板替换】完成')
print('储层表添加中，请等待...')

################################################################################
# 储层表的嵌入
PATH = "."
for fileName in os.listdir(PATH):
    newFile = PATH + "\\" + well_Name + '_' + year + month + \
              day + '_(' + casing_Goal + 'mm套,' + process_Section1 + 'm)固井报告' + '.docx'
document = Document(newFile)

if formation_be_or_not == '有储层':
    PATH = ".\\3储层表"
    for fileName in os.listdir(PATH):
        fileDir = PATH + "\\" + fileName
        workbook = xlrd.open_workbook(fileDir)

    sheet = workbook.sheets()[0]

    # 获得表单的行数及列数
    nrow = sheet.nrows
    ncol = sheet.ncols

    formation_table = document.tables[4]
    for num in range(eval(formation_Number) - 1):
        row_cells = formation_table.add_row()

    for row in range(1, len(formation_table.rows)):
        # print('已添加第', str(row), '个储层')
        for col in range(len(formation_table.columns)):
            formation_table.cell(row, col).text = str(sheet.cell_value(row + 1, col)).strip()
            formation_table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
            # print(formation_table.cell(row, col).text)
            formation_table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 首列居中
    for row in range(len(formation_table.rows)):
        formation_table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    print('【储层表】写入完成')
    document.save(newFile)  # 保存下

################################################################################
print('单层统计表添加中...')
PATH = ".\\3单层统计表"
for fileName in os.listdir(PATH):
    if '1单' in fileName and '.xlsx' in fileName and '$' not in fileName:
        fileDir = PATH + "\\" + fileName
        workbook = xlrd.open_workbook(fileDir)
sheet = workbook.sheets()[0]

# 通过xlrd的接口获得表单的行数及列数
nrow = sheet.nrows
ncol = sheet.ncols

document = Document(newFile)
document.styles['Normal'].font.size = Pt(9)  # 小五
document.styles['Normal'].font.name = u'Times New Roman'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 不采用堆砌方案，转为模板方案
# table = document.add_table(rows=nrow,cols=ncol,style='Table Grid')

if formation_be_or_not == '有储层':
    table = document.tables[6]
else:
    table = document.tables[5]
table.autofit = True
for num in range(nrow - 2):
    row_cells = table.add_row()

# 设置整个表格字体属性
table.style.font.color.rgb = RGBColor(0, 0, 0)
table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table.cell(0, 0).width = Pt(30)
table.cell(0, 1).width = Pt(100)

# 单层评价表写入
for row in range(len(table.rows)):
    # print('已添加第', str(row + 1), '行单层评价')
    view_bar(row, len(table.rows) - 1)
    for col in range(len(table.columns)):
        table.cell(row, col).text = str(sheet.cell_value(row, col))

# set_or_not = input('\n是否进行表格自动格式调整(耗时约1~3分钟)？\n【自动调整】请按【1】,【手动更改】请按【其它任意键】')
set_or_not = gui.indexbox(msg="是否进行表格自动格式调整(耗时约1-3分钟)？", title="提示", choices=("自动调整", "手动调整"))
if set_or_not == 0:
    # 合并单元格
    print('\n单层统计表合并单元格并居中...')
    for row in range(len(table.rows)):
        view_bar(row, len(table.rows) - 1)
        table.rows[row].height = Pt(20)
        for col in range(len(table.columns)):
            table.cell(row, 0).text = str(row - 1)  # 因为序号带小数，重新赋值
            table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    table.cell(0, 0).text = '序号'
    table.cell(1, 0).text = ''
    # 首列居中
    for row in range(len(table.rows)):
        table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 前两行居中
    for col in range(len(table.columns)):
        table.cell(0, col).merge(table.cell(1, col))

print('\n【一界面单层统计表】添加完成')


print('正在添加储层段落，请等待...')
################################################################################
# 上部井段固井质量评价表单元格居左
upper_interval_table = document.tables[3]
for row in range(1, len(upper_interval_table.rows)):
    for col in range(len(upper_interval_table.columns)):
        upper_interval_table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

################################################################################
# 储层固井质量评价
p = document.add_paragraph()
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
p.paragraph_format.line_spacing = Pt(24)
run = p.add_run(u"3.储层段固井质量分析")
run.font.name = 'Times New Roman'  # 英文字体
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)

evaluation_of_formation_upper = []
all_evaluation_of_formation_upper = []
if formation_be_or_not == '有储层':
    # 创建一个空的Dataframe
    formation_pic_DataFrame = pd.DataFrame(columns=('formation_StartNumber', 'formation_EndNumber', \
                                                    'formation_Start_Depth', 'formation_End_Depth'))
    # 添加段落
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
    r = p.add_run('该次测井井段有' + str(formation_Number) + '个解释储层。')
    # r.bold = True
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)

    PATH = ".\\4储层图"
    # 储层图片名
    all_Formation_Names = []
    # 储层图片后缀
    all_Formation_Extentions = []
    # 图片的数量
    count = 0
    for fileName in os.listdir(PATH):
        count += 1
        all_Formation_Names.append(get_filePath_fileName_fileExt(fileName)[1])
        all_Formation_Extentions.append(get_filePath_fileName_fileExt(fileName)[2])
    # 利用lambda表达式排序
    all_Formation_Names.sort(key=lambda x: int(x.split('#')[0].split('-')[0]))
    ################################################################################
    # 得到储层上部固井质量评价深度DataFrame
    for pic_number in range(count):
        formation_Name_Split = all_Formation_Names[pic_number].split('#')
        if '-' in formation_Name_Split[0]:
            formation_StartNumber = formation_Name_Split[0].split('-')[0]
            formation_EndNumber = formation_Name_Split[0].split('-')[1]
            formation_Start_End_Number = ''.join([formation_StartNumber, '-', formation_EndNumber])
        else:
            formation_StartNumber = formation_Name_Split[0]
            formation_EndNumber = formation_Name_Split[0]
            formation_Start_End_Number = formation_Name_Split[0]
        formation_Start_Depth = formation_Name_Split[1].split('-')[0]
        formation_End_Depth = formation_Name_Split[1].split('-')[1]
        formation_Start_End = ''.join([formation_Start_Depth, '-', formation_End_Depth])
        ################################################################################
        # 重复代码2，待精简
        # 单层统计表
        PATH = ".\\3单层统计表"
        for fileName in os.listdir(PATH):
            fileDir = PATH + "\\" + fileName
        df1 = pd.read_excel(fileDir, header=0, index='序号')
        df1.drop([0], inplace=True)
        df1.drop(['平均声幅', '最小声幅', '最大声幅'], axis=1, inplace=True)
        df1.loc[:, '井段'] = df1['井段'].str.replace(' ', '')  # 消除数据中空格
        df1['井段Start'] = df1['井段'].map(lambda x: x.split("--")[0])
        df1['井段End'] = df1['井段'].map(lambda x: x.split("--")[1])

        # 储层表
        PATH = ".\\3储层表"
        for fileName in os.listdir(PATH):
            fileDir = PATH + "\\" + fileName
        df2 = pd.read_excel(fileDir, header=0, index='序号')
        df2.drop([0], inplace=True)
        df2.drop(['层位', '解释结论'], axis=1, inplace=True)
        df2.loc[:, '井        段'] = df2['井        段'].str.replace(' ', '')  # 消除数据中空格
        df2['储层Start'] = df2['井        段'].map(lambda x: x.split("--")[0])
        df2['储层End'] = df2['井        段'].map(lambda x: x.split("--")[1])

        # 当前储层图片里第一个储层的上界深度
        formation_Start = df2.loc[int(formation_StartNumber), '储层Start']
        # 添加要分析的DataFrame
        formation_pic_DataFrame = formation_pic_DataFrame.append(
            pd.DataFrame(
                {'formation_StartNumber': [formation_StartNumber], 'formation_EndNumber': [formation_EndNumber], \
                 '当前储层图片Start': [formation_Start_Depth], '第一个储层start': [formation_Start]}), \
            ignore_index=True)
    # print(formation_pic_DataFrame)

    ################################################################################
    # 重复代码3，待精简
    # 单层统计表
    PATH = ".\\3单层统计表"
    for fileName in os.listdir(PATH):
        fileDir = PATH + "\\" + fileName
    df1 = pd.read_excel(fileDir, header=0, index='序号')
    df1.drop([0], inplace=True)
    df1.drop(['平均声幅', '最小声幅', '最大声幅'], axis=1, inplace=True)
    df1.loc[:, '井段'] = df1['井段'].str.replace(' ', '')  # 消除数据中空格
    df1['井段Start'] = df1['井段'].map(lambda x: x.split("--")[0])
    df1['井段End'] = df1['井段'].map(lambda x: x.split("--")[1])

    # 储层表
    PATH = ".\\3储层表"
    for fileName in os.listdir(PATH):
        fileDir = PATH + "\\" + fileName
    df2 = pd.read_excel(fileDir, header=0, index='序号')
    df2.drop([0], inplace=True)
    df2.drop(['层位', '解释结论'], axis=1, inplace=True)
    df2.loc[:, '井        段'] = df2['井        段'].str.replace(' ', '')  # 消除数据中空格
    df2['储层Start'] = df2['井        段'].map(lambda x: x.split("--")[0])
    df2['储层End'] = df2['井        段'].map(lambda x: x.split("--")[1])
    # 表格数据清洗
    df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
    df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
    df2.loc[:, "储层Start"] = df2["储层Start"].str.replace(" ", "").astype('float')
    df2.loc[:, "储层End"] = df2["储层End"].str.replace(" ", "").astype('float')
    rows1, cols1 = df1.shape
    rows2, cols2 = df2.shape

    # 针对每层在单层评价表中得出好中差比例
    for pic_number in range(count):
        calculation_Start = formation_pic_DataFrame.loc[pic_number, '当前储层图片Start']
        calculation_Start = float(calculation_Start)
        calculation_End = formation_pic_DataFrame.loc[pic_number, '第一个储层start']
        calculation_End = float(calculation_End)
        # print('----------------第', pic_number + 1, '个储层上部井段----------------')
        if (calculation_End <= float(end_Evaluation)) & (calculation_Start >= float(start_Evaluation)):
            df_temp = df1.loc[(df1['井段Start'] >= calculation_Start) & (df1['井段Start'] <= calculation_End), :]
            # 获取储层起始深度到第一层井段底界的结论
            df_temp1 = df1.loc[(df1['井段Start'] <= calculation_Start), :]
            start_to_upper_result = df_temp1.loc[len(df_temp1), '结论']
            # 补充储层界到井段的深度
            x, y = df_temp.shape
            df_temp = df_temp.reset_index()
            df_temp.drop(['index'], axis=1, inplace=True)
            if x >= 1:
                first_layer_start = df_temp.loc[0, '井段Start']
                upper = pd.DataFrame({'序号': '空',
                                      '井段': '空',
                                      '厚度': '空',
                                      '结论': start_to_upper_result,
                                      '井段Start': calculation_Start,
                                      '井段End': first_layer_start},
                                     index=[1])  # 自定义索引为：1 ，这里也可以不设置index
                df_temp.loc[len(df_temp) - 1, '井段End'] = calculation_End
                df_temp = pd.concat([upper, df_temp], ignore_index=True)
                # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
                # print(df_temp)
            else:  # 储层包含在一个井段内的情况
                df_temp = pd.DataFrame({'序号': '空',
                                        '井段': '空',
                                        '厚度': '空',
                                        '结论': start_to_upper_result,
                                        '井段Start': calculation_Start,
                                        '井段End': calculation_End},
                                       index=[1])  # 自定义索引为：1 ，这里也可以不设置index
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(get_thickness, axis=1)
            # print(df_temp)
            ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100
            if ratio_Series.__len__() == 2:
                if '好' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
                elif '中' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
                elif '差' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'差': 0}))
            elif ratio_Series.__len__() == 1:
                if ('好' not in ratio_Series) & ('中' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
                elif ('好' not in ratio_Series) & ('差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'差': 0}))
                elif ('中' not in ratio_Series) & ('差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'中': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'差': 0}))

            # 条件判断，参数需要研究
            if ratio_Series['好'] >= 95:
                evaluation_of_formation_upper = '好'
            elif ratio_Series['中'] >= 95:
                evaluation_of_formation_upper = '中'
            elif ratio_Series['差'] >= 95:
                evaluation_of_formation_upper = '差'
            elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['中'] >= 5) & (5 >= ratio_Series['差']):
                if ratio_Series['好'] >= ratio_Series['中']:
                    evaluation_of_formation_upper = '中到好，以好为主'
                elif ratio_Series['好'] <= ratio_Series['中']:
                    evaluation_of_formation_upper = '中到好，以中等为主'
            elif (95 >= ratio_Series['差'] >= 5) & (95 >= ratio_Series['中'] >= 5) & (5 >= ratio_Series['好']):
                if ratio_Series['差'] >= ratio_Series['中']:
                    evaluation_of_formation_upper = '中到差，以差为主'
                elif ratio_Series['差'] <= ratio_Series['中']:
                    evaluation_of_formation_upper = '中到差，以中等为主'
            elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['差'] >= 5) & (5 >= ratio_Series['中']):
                if ratio_Series['好'] >= ratio_Series['差']:
                    evaluation_of_formation_upper = '好到差，以好为主'
                elif ratio_Series['好'] <= ratio_Series['差']:
                    evaluation_of_formation_upper = '好到差，以差为主'
            elif (95 >= ratio_Series['好'] >= 5) & (95 >= ratio_Series['差'] >= 5) & (95 >= ratio_Series['中'] >= 5):
                evaluation_of_formation_upper = '好到中到差'
            # print(ratio_Series)  # 某一储层的评价
            # print(evaluation_of_formation_upper)  # 全部储层的描述list
            all_evaluation_of_formation_upper.append(evaluation_of_formation_upper)
        else:
            # print('统计范围超出了测量范围，请检查')
            pass
    # print(all_evaluation_of_formation_upper)
    ################################################################################
    # 储层上部描述输出
    # 图片的数量
    PATH = ".\\4储层图"
    count = 0
    for fileName in os.listdir(PATH):
        count += 1
    for pic_number in range(count):
        formation_Name_Split = all_Formation_Names[pic_number].split('#')
        if '-' in formation_Name_Split[0]:
            formation_StartNumber = formation_Name_Split[0].split('-')[0]
            formation_EndNumber = formation_Name_Split[0].split('-')[1]
            formation_Start_End_Number = ''.join([formation_StartNumber, '-', formation_EndNumber])
        else:
            formation_StartNumber = formation_Name_Split[0]
            formation_EndNumber = formation_Name_Split[0]
            formation_Start_End_Number = formation_Name_Split[0]
        formation_Start_Depth = formation_Name_Split[1].split('-')[0]
        formation_End_Depth = formation_Name_Split[1].split('-')[1]
        formation_Start_End = ''.join([formation_Start_Depth, '-', formation_End_Depth])
        ###
        if all_evaluation_of_formation_upper[pic_number] == '好':
            cbl_value = '低'
            case_Wave_Energy = '弱'
            formation_Wave_Energy = '强'
        elif all_evaluation_of_formation_upper[pic_number] == '中':
            cbl_value = '中'
            case_Wave_Energy = '较强'
            formation_Wave_Energy = '较弱'
        elif all_evaluation_of_formation_upper[pic_number] == '差':
            cbl_value = '高'
            case_Wave_Energy = '强'
            formation_Wave_Energy = '弱'
        elif all_evaluation_of_formation_upper[pic_number] in ['中到好，以好为主', '中到好，以中等为主']:
            cbl_value = '中到低'
            case_Wave_Energy = '较强到较弱'
            formation_Wave_Energy = '较弱到较强'
        elif all_evaluation_of_formation_upper[pic_number] in ['中到差，以差为主', '中到差，以中等为主']:
            cbl_value = '中到高'
            case_Wave_Energy = '较强到强'
            formation_Wave_Energy = '较弱到弱'
        elif all_evaluation_of_formation_upper[pic_number] in ['好到差，以好为主', '好到差，以差为主']:
            cbl_value = '低到高'
            case_Wave_Energy = '强到弱'
            formation_Wave_Energy = '弱到强'
        elif all_evaluation_of_formation_upper[pic_number] in ['好到中到差']:
            cbl_value = '低到高'
            case_Wave_Energy = '较强到较弱'
            formation_Wave_Energy = '较弱到较强'
        ###
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
        r = p.add_run('（' + str(pic_number + 1) + '）' + formation_Start_End + 'm该封固井段上部声幅值')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        r = p.add_run(cbl_value)
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 250)

        r = p.add_run('，一界面水泥胶结')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        r = p.add_run(all_evaluation_of_formation_upper[pic_number])
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 250)

        r = p.add_run('；变密度曲线反映内层套管波能量')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        r = p.add_run(case_Wave_Energy)
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 250)

        r = p.add_run('，地层波能量')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        r = p.add_run(formation_Wave_Energy)
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 250)

        r = p.add_run('，二界面水泥胶结')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        r = p.add_run('[待确定]')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 0, 0)

        r = p.add_run('。')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        ########################################################################################
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符

        # 储层描述循环输出
        for formation_Number_Temp in range(int(formation_StartNumber), int(formation_EndNumber) + 1):
            if formation_Number_Temp <= len(all_evaluation_of_formation):
                if all_evaluation_of_formation[formation_Number_Temp - 1] == '好':
                    cbl_value = '低'
                    case_Wave_Energy = '弱'
                    formation_Wave_Energy = '强'
                elif all_evaluation_of_formation[formation_Number_Temp - 1] == '中':
                    cbl_value = '中'
                    case_Wave_Energy = '较强'
                    formation_Wave_Energy = '较弱'
                elif all_evaluation_of_formation[formation_Number_Temp - 1] == '差':
                    cbl_value = '高'
                    case_Wave_Energy = '强'
                    formation_Wave_Energy = '弱'
                elif all_evaluation_of_formation[formation_Number_Temp - 1] in ['中到好，以好为主', '中到好，以中等为主']:
                    cbl_value = '中到低'
                    case_Wave_Energy = '较强到较弱'
                    formation_Wave_Energy = '较弱到较强'
                elif all_evaluation_of_formation[formation_Number_Temp - 1] in ['中到差，以差为主', '中到差，以中等为主']:
                    cbl_value = '中到高'
                    case_Wave_Energy = '较强到强'
                    formation_Wave_Energy = '较弱到弱'
                elif all_evaluation_of_formation[formation_Number_Temp - 1] in ['好到差，以好为主', '好到差，以差为主']:
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到弱'
                    formation_Wave_Energy = '弱到强'
                elif all_evaluation_of_formation[formation_Number_Temp - 1] in ['好到中到差']:
                    cbl_value = '低到高'
                    case_Wave_Energy = '较强到较弱'
                    formation_Wave_Energy = '较弱到较强'
            else:
                all_evaluation_of_formation.append('[储层范围超出测量边界，待确定]')

            if all_evaluation_of_formation[formation_Number_Temp - 1] == '[储层范围超出测量边界，待确定]':
                cbl_value = '[储层范围超出测量边界，待确定]'
                case_Wave_Energy = '[储层范围超出测量边界，待确定]'
                formation_Wave_Energy = '[储层范围超出测量边界，待确定]'
            r = p.add_run(str(formation_Number_Temp) + '#储层声幅')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            r = p.add_run(cbl_value)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 250)

            r = p.add_run('，一界面水泥胶结')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            r = p.add_run(all_evaluation_of_formation[formation_Number_Temp - 1])
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 250)

            r = p.add_run('；变密度曲线反映套管波能量')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            r = p.add_run(case_Wave_Energy)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 250)

            r = p.add_run('、地层波能量')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            r = p.add_run(formation_Wave_Energy)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 250)

            r = p.add_run('，二界面水泥胶结')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            r = p.add_run('[待确定]。')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(250, 0, 0)

        r = p.add_run('（见图' + str(pic_number + 1) + '）')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        # 添加储层图片
        paragraph = document.add_paragraph()
        # 图片居中设置
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run("")
        PATH = ".\\4储层图"
        if all_Formation_Extentions[0] == '.png':
            run.add_picture(PATH + '\\' + all_Formation_Names[pic_number] + '.png', width=Inches(5.0))
        elif all_Formation_Extentions[0] == '.jpg':
            run.add_picture(PATH + '\\' + all_Formation_Names[pic_number] + '.jpg', width=Inches(5.0))
        elif all_Formation_Extentions[0] == '.bmp':
            run.add_picture(PATH + '\\' + all_Formation_Names[pic_number] + '.bmp', width=Inches(5.0))
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = Pt(24)
        r = p.add_run('图' + str(pic_number + 1) + '  ' + well_Name + '井（' + formation_Start_End + 'm）固井处理成果图')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        # print('已添加第', str(pic_number + 1), '个储层的段落')
else:
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
    r = p.add_run('该次测量井段内无储层解释。')
    # r.bold = True
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)
################################################################################
# 判断是否有固井为差
PATH = ".\\5胶结差图\\"
bad_interval_be_or_not = ''
if os.listdir(PATH) != []:
    for fileName in os.listdir(PATH):
        fileName = fileName
else:
    fileName = ''

f_path = PATH + fileName

if os.path.isdir(f_path):
    bad_interval_be_or_not = '无固井差'
else:
    bad_interval_be_or_not = '有固井差'
################################################################################
if bad_interval_be_or_not == '有固井差':
    PATH = ".\\5胶结差图"
    bad_Interval_Names = []
    bad_Interval_Extentions = []
    for fileName in os.listdir(PATH):
        bad_Interval_Names.append(get_filePath_fileName_fileExt(fileName)[1])
        bad_Interval_Extentions.append(get_filePath_fileName_fileExt(fileName)[2])

    bad_Start_Ends = []
    for bad_number in range(len(bad_Interval_Names)):
        bad_Name_Split = bad_Interval_Names[bad_number].split('-')
        bad_Serial_Number = bad_Name_Split[0]
        bad_Start_Depth = bad_Name_Split[1]
        bad_End_Depth = bad_Name_Split[2]
        bad_Start_End = ''.join([bad_Start_Depth, '-', bad_End_Depth])
        bad_Start_Ends.append(bad_Start_End + 'm、')
    bad_Start_Ends = ''.join(bad_Start_Ends).rstrip('、')

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(u"三 建议及其它")
    run.font.name = '黑体'  # 英文字体
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 中文字体
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    if formation_be_or_not == '无储层':
        pic_number = -1
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
    if bad_number != 0:
        r = p.add_run(
            bad_Start_Ends + '井段声幅值较高，部分套管接箍信号明显，建议采取相应措施（见图' + str(pic_number + 2) + '-' + str(
                pic_number + bad_number + 2) + '）。')
    elif bad_number == 0:
        r = p.add_run(
            bad_Start_Ends + '井段声幅值较高，部分套管接箍信号明显，建议采取相应措施（见图' + str(pic_number + 2) + '）。')
    # r.bold = True
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)

    print('正在添加固井为差段落，请等待...')
    # 添加固井质量差图片
    for bad_number in range(len(bad_Interval_Names)):
        bad_Name_Split = bad_Interval_Names[bad_number].split('-')
        bad_Serial_Number = bad_Name_Split[0]
        bad_Start_Depth = bad_Name_Split[1]
        bad_End_Depth = bad_Name_Split[2]
        bad_Start_End = ''.join([bad_Start_Depth, '-', bad_End_Depth])

        paragraph = document.add_paragraph()
        # 图片居中设置
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run("")

        if bad_Interval_Extentions[0] == '.png':
            run.add_picture(PATH + '\\' + bad_Interval_Names[bad_number] + '.png', width=Inches(5.0))
        elif bad_Interval_Extentions[0] == '.jpg':
            run.add_picture(PATH + '\\' + bad_Interval_Names[bad_number] + '.jpg', width=Inches(5.0))
        elif bad_Interval_Extentions[0] == '.bmp':
            run.add_picture(PATH + '\\' + bad_Interval_Names[bad_number] + '.bmp', width=Inches(5.0))

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = Pt(24)
        r = p.add_run('图' + str(pic_number + bad_number + 2) + ' ' + well_Name + '井（' + bad_Start_End + 'm）固井处理成果图')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        # print('已添加第', str(bad_number + 1), '个固井为差的段落')
else:
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(u"三 建议及其它")
    run.font.name = '黑体'  # 英文字体
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 中文字体
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
    r = p.add_run('测量井段内的固井质量以好为主。')
    # r.bold = True
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)
################################################################################
# 签名
PATH = '.\\resources\\签名图片\\'
choicess_list = ["李海军", "陈海祥", "朱莉", "何强", "杨艺", "罗文", "王昌德"]
report_Writer = gui.choicebox(msg='请选择报告编写人', choices=choicess_list)
add1 = document.tables[0].cell(0, 1).paragraphs[0]
if report_Writer == '李海军':
    add1.add_run().add_picture(PATH + '签名-李海军.jpg', width=Inches(1.0))
elif report_Writer == '陈海祥':
    add1.add_run().add_picture(PATH + '签名-陈海祥.jpg', width=Inches(1.0))
elif report_Writer == '朱莉':
    add1.add_run().add_picture(PATH + '签名-朱莉.jpg', width=Inches(1.0))
elif report_Writer == '何强':
    add1.add_run().add_picture(PATH + '签名-何强.jpg', width=Inches(1.0))
elif report_Writer == '杨艺':
    add1.add_run().add_picture(PATH + '签名-杨艺.jpg', width=Inches(1.0))
elif report_Writer == '罗文':
    add1.add_run().add_picture(PATH + '签名-罗文.jpg', width=Inches(1.0))
elif report_Writer == '王昌德':
    add1.add_run().add_picture(PATH + '签名-王昌德.jpg', width=Inches(1.0))
choicess_list = ["刘恒", "王参文", "刘静", "朱莉"]
report_Supervisor = gui.choicebox(msg='请选择报告审核人', choices=choicess_list)
add2 = document.tables[0].cell(1, 1).paragraphs[0]
if report_Supervisor == '刘恒':
    add2.add_run().add_picture(PATH + '签名-刘恒.jpg', width=Inches(1.0))
elif report_Supervisor == '王参文':
    add2.add_run().add_picture(PATH + '签名-王参文.jpg', width=Inches(1.0))
elif report_Supervisor == '刘静':
    add2.add_run().add_picture(PATH + '签名-刘静.jpg', width=Inches(1.0))
elif report_Supervisor == '朱莉':
    add2.add_run().add_picture(PATH + '签名-朱莉.jpg', width=Inches(1.0))
document.save(newFile)
print('【报告】生成完毕')
input('按任意键退出'.center(25, '-'))
